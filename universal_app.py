import streamlit as st
import asyncio
import nest_asyncio
nest_asyncio.apply()
import sys
import os
import json
import requests
import subprocess
import re
import random
from datetime import datetime
from io import BytesIO
from typing import List, Optional, Dict
from pydantic import BaseModel, Field
from playwright.async_api import async_playwright, Page, Browser, BrowserContext

# Fehlende Imports für den Word-Export hinzugefügt:
from docx import Document
from docx.shared import Cm, Pt
from PIL import Image

# Standard Windows Policy for Playwright
if sys.platform == 'win32':
    asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())

# --------------------------------------------------
# DATA MODELS
# --------------------------------------------------
class PropertyListing(BaseModel):
    name: str
    price: str
    location: str
    area: str
    image_url: Optional[str]
    link: str
    source: str

# --------------------------------------------------
# SCRAPER CORE (Integrated)
# --------------------------------------------------
class ScraperCore:
    def __init__(self):
        self.browser: Optional[Browser] = None
        self.context: Optional[BrowserContext] = None

    async def init_browser(self, headless: bool = True):
        if not self.browser:
            p = await async_playwright().start()
            self.browser = await p.chromium.launch(
                headless=headless,
                args=[
                    '--disable-blink-features=AutomationControlled',
                    '--no-sandbox',
                    '--disable-setuid-sandbox',
                    '--disable-infobars',
                    '--window-position=0,0',
                    '--ignore-certifcate-errors',
                    '--ignore-certifcate-errors-spki-list',
                    '--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36'
                ]
            )
            self.context = await self.browser.new_context(
                viewport={'width': 1280, 'height': 800},
                user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36",
                locale="de-DE",
                timezone_id="Europe/Berlin"
            )
            # Advanced Stealth Script
            await self.context.add_init_script("""
                Object.defineProperty(navigator, 'webdriver', { get: () => undefined });
                window.chrome = { runtime: {} };
                const originalQuery = window.navigator.permissions.query;
                window.navigator.permissions.query = (parameters) => (
                    parameters.name === 'notifications' ?
                        Promise.resolve({ state: Notification.permission }) :
                        originalQuery(parameters)
                );
                Object.defineProperty(navigator, 'plugins', { get: () => [1, 2, 3, 4, 5] });
                Object.defineProperty(navigator, 'languages', { get: () => ['de-DE', 'de', 'en-US', 'en'] });
            """)

    async def close_browser(self):
        if self.browser:
            await self.browser.close()
            self.browser = None

    async def _clean_modals(self, page: Page):
        try:
            # Handle Usercentrics specifically
            await page.evaluate("""() => {
                const uc = document.querySelector('#usercentrics-root');
                if (uc && uc.shadowRoot) {
                    const btn = uc.shadowRoot.querySelector('button[data-testid="uc-accept-all-button"]');
                    if (btn) btn.click();
                }
            }""")
            await asyncio.sleep(1)
            # Remove generic banners
            await page.evaluate("""
                const selectors = [
                    '#usercentrics-cmp', '#usercentrics-root', '.nlf-slidebox', 
                    '.modal', '.cookie-consent', '.banner', '#onetrust-consent-sdk',
                    '.cky-consent-container', '.cc-window', '#c-p-bn'
                ];
                selectors.forEach(s => {
                    const el = document.querySelector(s);
                    if (el) el.remove();
                });
                document.body.style.overflow = 'auto';
                document.documentElement.style.overflow = 'auto';
            """)
        except: pass

    async def scrape_site(self, url: str, headless: bool = True) -> List[PropertyListing]:
        await self.init_browser(headless=headless)
        page = await self.context.new_page()
        data = []
        try:
            print(f"Scraping: {url}")
            await page.goto(url, wait_until="load", timeout=60000)
            await page.wait_for_timeout(5000)
            
            url_lower = url.lower()
            if "homebase" in url_lower: data = await self.parse_homebase(page)
            elif "kensington" in url_lower: data = await self.parse_kensington(page)
            elif "teampower" in url_lower: data = await self.parse_teampower(page)
            elif "deutsches-immowerk" in url_lower or "immowerk" in url_lower: data = await self.parse_immowerk(page)
            elif "robertcspies" in url_lower: data = await self.parse_robertcspies(page)
            elif "pump" in url_lower: data = await self.parse_pump(page)
            else: data = []
        except Exception as e:
            print(f"Error scraping {url}: {e}")
        finally:
            await page.close()
        return data

    async def parse_kensington(self, page: Page):
        all_listings = []
        await page.wait_for_timeout(3000)
        await self._clean_modals(page)
        
        # Check Cloudflare
        content = await page.content()
        if any(x in content for x in ["Just a moment", "Verify you are human", "cf-challenge"]):
            print("Kensington: CLOUDFLARE! Warten...")
            try: await page.wait_for_selector(".card", timeout=150000); await self._clean_modals(page)
            except: return []

        current_page = 1
        max_pages = 9
        while current_page <= max_pages:
            print(f"Kensington Seite {current_page}: Scrolle...")
            for i in range(5):
                await page.evaluate(f"window.scrollBy(0, 800)"); await page.wait_for_timeout(700)
            
            items = page.locator(".card")
            count = await items.count()
            if count == 0: break

            for i in range(count):
                try:
                    it = items.nth(i)
                    title_el = it.locator("h5.title a").first
                    if await title_el.count() == 0: continue
                    name = (await title_el.inner_text()).strip()
                    if not name: continue
                    
                    full_text = await it.inner_text()
                    price = "n.v."
                    price_el = it.locator(".gold p, .gold.mb-3 p").first
                    if await price_el.count() > 0: price = (await price_el.inner_text()).strip()
                    if price == "n.v.":
                        m = re.search(r"[\d.,]+\s*(?:EUR|€)", full_text)
                        if m: price = m.group(0)

                    area = "n.v."
                    area_els = it.locator(".werte")
                    for j in range(await area_els.count()):
                        t = await area_els.nth(j).inner_text()
                        if "m²" in t: area = t.strip(); break
                    
                    location = "Hamburg"
                    loc_el = it.locator(".card-body .small.gold").first
                    if await loc_el.count() > 0: location = (await loc_el.inner_text()).strip()

                    img_el = it.locator("img").first
                    image_url = ""
                    if await img_el.count() > 0:
                        image_url = await img_el.evaluate("el => el.src || el.dataset.src || el.dataset.lazySrc || ''")
                    
                    link = await title_el.get_attribute("href") or ""
                    if link and not link.startswith("http"): link = "https://kensington-international.com" + link

                    all_listings.append(PropertyListing(
                        name=name, price=price, location=location, area=area,
                        image_url=str(image_url) if image_url else "", link=link, source="Kensington"
                    ))
                except: continue

            next_btn = page.locator(".pagination li.page-item:not(.active):not(.disabled) a.page-link").last
            if await next_btn.count() > 0 and await next_btn.is_visible():
                print(f"Kensington: Weiter zu Seite {current_page+1}...")
                await next_btn.click(); await page.wait_for_timeout(5000); current_page += 1
            else: break
        return all_listings

    async def parse_homebase(self, page: Page):
        all_listings = []
        max_pages = 4
        page_num = 1
        while page_num <= max_pages:
            print(f"Homebase Seite {page_num}...")
            try: await page.locator(".obj-list-object").first.wait_for(timeout=15000)
            except: break

            await page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
            await page.wait_for_timeout(2000)

            loc = page.locator(".obj-list-object")
            count = await loc.count()
            for i in range(count):
                try:
                    it = loc.nth(i)
                    name = await it.locator(".obj-list-title span, h3").first.inner_text()
                    price_el = it.locator(".obj-kaufpreis, .obj-price").first
                    price = await price_el.inner_text() if await price_el.count() > 0 else "n.v."
                    loc_el = it.locator(".obj-geo span").first
                    location = await loc_el.inner_text() if await loc_el.count() > 0 else "Hamburg"
                    
                    area = "n.v."
                    area_el = it.locator(".object-area-value, .obj-list-data").first
                    if await area_el.count() > 0:
                        t = await area_el.inner_text()
                        m = re.search(r"(\d+\s*m²)", t)
                        if m: area = m.group(1)

                    img_el = it.locator("img").first
                    img = ""
                    if await img_el.count() > 0:
                        img = await img_el.evaluate("el => el.src || el.dataset.src || el.dataset.lazySrc || ''")
                    
                    link_el = it.locator("h5.title a, a[href*='/objekte/'], a[href*='/expose/']").first
                    link = await link_el.get_attribute("href") if await link_el.count() > 0 else ""
                    if link and not link.startswith("http"): link = "https://www.homebase-immobilienberatung.de" + link

                    all_listings.append(PropertyListing(
                        name=name.strip(), price=price.strip(), location=location.strip(),
                        area=area.strip(), image_url=str(img) if img else "", link=link, source="Homebase"
                    ))
                except: continue

            if page_num < max_pages:
                next_btn = page.locator(f"a[href*='p[obj0]={page_num + 1}']").first
                if await next_btn.count() == 0:
                    next_btn = page.locator("a i.fa-angle-right, a i.fas.fa-angle-right").locator("xpath=..").first
                if await next_btn.count() > 0:
                    await next_btn.click(); await page.wait_for_timeout(6000); page_num += 1
                else: break
            else: break
        return all_listings

    async def parse_teampower(self, page: Page):
        try:
            btn = await page.query_selector(".geowerft-switcher-left")
            if btn: await btn.click(); await asyncio.sleep(2)
        except: pass
        items = await page.query_selector_all(".single-object")
        listings = []
        for it in items:
            try:
                name = await (await it.query_selector(".object-titel")).inner_text()
                price = await (await it.query_selector(".object-pricing")).inner_text()
                location = await (await it.query_selector(".city")).inner_text()
                link = await (await it.query_selector(".geowerft-list-anker")).get_attribute("href")
                img_el = await it.query_selector(".mw-geowerft-list-single-prop-img")
                img = await img_el.evaluate("el => el.src || ''") if img_el else ""
                
                listings.append(PropertyListing(
                    name=name.strip(), price=price.strip(), location=location.strip(), 
                    area=self._extract_area(await it.inner_text()), image_url=img, link=link, source="TeamPower"
                ))
            except: continue
        return listings

    async def parse_immowerk(self, page: Page):
        items_data = await page.evaluate("""() => {
            const h2s = Array.from(document.querySelectorAll('h2'));
            const successHeader = h2s.find(h => h.innerText.includes('EINE AUSWAHL UNSERER VERKAUFSERFOLGE'));
            const allItems = Array.from(document.querySelectorAll('.elementor-widget-wrap'));
            let validItems = [];
            for (let item of allItems) {
                if (successHeader && (item.compareDocumentPosition(successHeader) & Node.DOCUMENT_POSITION_PRECEDING)) break;
                if (item.querySelector('.elementor-cta__title')) {
                    validItems.push({
                        name: item.querySelector('.elementor-cta__title').innerText.trim(),
                        price: item.querySelector('.elementor-cta__description') ? item.querySelector('.elementor-cta__description').innerText.trim() : 'n.v.',
                        link: item.querySelector('.elementor-cta__button') ? item.querySelector('.elementor-cta__button').href : window.location.href,
                        img: item.querySelector('.elementor-widget-image img') ? item.querySelector('.elementor-widget-image img').src : null
                    });
                }
            }
            return validItems;
        }""")
        listings = [PropertyListing(name=d['name'], price=d['price'], location="Hamburg", area=self._extract_area(d['name']), image_url=d['img'], link=d['link'], source="Immowerk") for d in items_data]
        return listings

    async def parse_robertcspies(self, page: Page):
        items = await page.query_selector_all(".exposeList__item")
        listings = []
        for it in items:
            try:
                name = await (await it.query_selector(".exposeList__item__content__linkedHeadline h2")).inner_text()
                link = await (await it.query_selector(".exposeList__item__content__linkedHeadline")).get_attribute("href")
                if link and not link.startswith("http"): link = "https://robertcspies.de" + link
                p, a = "n.v.", "n.v."
                for f in await it.query_selector_all("li.field"):
                    l = await (await f.query_selector(".label")).inner_text()
                    v = await (await f.query_selector(".value")).inner_text()
                    if "Kaufpreis" in l: p = v
                    if "Wohnfläche" in l: a = v
                img_el = await it.query_selector("img")
                img = await img_el.evaluate("el => el.src || ''") if img_el else ""
                listings.append(PropertyListing(name=name.strip(), price=p.strip(), location="Hamburg", area=a.strip(), image_url=img, link=link, source="Robert C. Spies"))
            except: continue
        return listings

    async def parse_pump(self, page: Page):
        items = await page.query_selector_all(".obj-list-object")
        listings = []
        for it in items:
            try:
                name = await (await it.query_selector("h2")).inner_text()
                link = await (await it.query_selector("a")).get_attribute("href")
                if link and not link.startswith("http"): link = "https://www.pump-immobilien.de/" + link
                price = await (await it.query_selector("p > span > span")).inner_text()
                area = await (await it.query_selector(".object-area-value")).inner_text()
                img_el = await it.query_selector("img")
                img = await img_el.evaluate("el => el.src || ''") if img_el else ""
                listings.append(PropertyListing(name=name.strip(), price=price.strip(), location="Hamburg", area=area.strip(), image_url=img, link=link, source="Pump Immobilien"))
            except: continue
        return listings

    def _extract_area(self, text: str):
        m = re.search(r"(\d+([.,]\d+)?\s*(m2|qm|m²))", text, re.IGNORECASE)
        return m.group(0) if m else "n.v."

    async def scrape_all(self, selected_sources: Optional[List[str]] = None, headless: bool = True):
        urls = {
            "Homebase": "https://www.homebase-immobilienberatung.de/suchergebnisse.xhtml?f%5B84369-9%5D=kauf",
            "Kensington": "https://kensington-international.com/de/de/hamburg/suche",
            "TeamPower": "https://teampower-immobilien.de/immobilien/",
            "Immowerk": "https://www.deutsches-immowerk.de/immobilien/",
            "Robert C. Spies": "https://robertcspies.de/wohnen/objekte?type=sale&region%5B%5D=4&price_max=850000",
            "Pump Immobilien": "https://www.pump-immobilien.de/immobilien.xhtml"
        }
        targets = {k: v for k, v in urls.items() if k in selected_sources} if selected_sources else urls
        all_data = []
        for name, url in targets.items():
            results = await self.scrape_site(url, headless=headless)
            all_data.extend(results)
            print(f"[{name}] Found {len(results)}.")
        return all_data

# --------------------------------------------------
# UI & EXPORT LOGIC
# --------------------------------------------------
HISTORY_FILE = "export_history.json"

def load_history():
    if os.path.exists(HISTORY_FILE):
        try:
            with open(HISTORY_FILE, "r") as f: return json.load(f)
        except: return {}
    return {}

def save_to_history(cart_items):
    history = load_history()
    now = datetime.now().strftime("%d.%m.%Y %H:%M")
    for link in cart_items.keys():
        if link: history[link] = now
    with open(HISTORY_FILE, "w") as f: json.dump(history, f, indent=4)

def set_compact(paragraph):
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

# CLOUD FIX: Schreibt in den Arbeitsspeicher statt auf den Desktop
def create_export_buffer():
    if not st.session_state.cart:
        return None
    
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
    title = doc.add_heading("Immobilien Recherche Bericht", 0)
    set_compact(title)
    
    for uid, it in st.session_state.cart.items():
        table = doc.add_table(rows=1, cols=2)
        table.columns[0].width = Cm(5.5)
        table.columns[1].width = Cm(10.5)
        
        if it.image_url:
            try:
                print(f"Versuche Bild-Download: {it.image_url}")
                headers = {
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36",
                    "Referer": "https://kensington-international.com/",
                    "Accept": "image/avif,image/webp,image/apng,image/*,*/*;q=0.8"
                }
                r = requests.get(it.image_url, timeout=15, headers=headers)
                r.raise_for_status()
                
                img = Image.open(BytesIO(r.content)).convert("RGB")
                buf = BytesIO()
                img.save(buf, format="JPEG", quality=85)
                buf.seek(0)
                pic_p = table.rows[0].cells[0].paragraphs[0]
                pic_p.add_run().add_picture(buf, width=Cm(5))
                set_compact(pic_p)
            except Exception as e:
                table.rows[0].cells[0].text = f"Bild n.v.\n({str(e)[:20]})"
        else:
            table.rows[0].cells[0].text = "Kein Bild"
        
        c2 = table.rows[0].cells[1]
        p_name = c2.paragraphs[0]
        n_run = p_name.add_run(it.name); n_run.bold = True; n_run.font.size = Pt(11); set_compact(p_name)
        for label, val in [("Ort", it.location), ("Preis", it.price), ("Fläche", it.area)]:
            p = c2.add_paragraph(f"{label}: {val}"); set_compact(p)
        p_link = c2.add_paragraph("Link: "); r_link = p_link.add_run(it.link); r_link.underline = True; r_link.font.size = Pt(9); set_compact(p_link)
        doc.add_paragraph("-" * 85)

    # Speichern im RAM-Buffer
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    save_to_history(st.session_state.cart)
    return file_stream.getvalue()

def toggle_item(link, it):
    if link in st.session_state.cart: del st.session_state.cart[link]
    else: st.session_state.cart[link] = it

# --------------------------------------------------
# STREAMLIT UI
# --------------------------------------------------
st.set_page_config(page_title="Immo Scraper 2026", layout="wide")

# CLOUD FIX: Playwright Browser Installation
if "playwright_installed" not in st.session_state:
    try:
        # Check if already installed
        subprocess.run(["playwright", "install", "chromium"], check=True, capture_output=True)
        st.session_state.playwright_installed = True
    except:
        with st.spinner("Initialisiere Browser-Umgebung (Playwright install)..."):
            subprocess.run([sys.executable, "-m", "playwright", "install", "chromium"])
            st.session_state.playwright_installed = True

if "listings" not in st.session_state: st.session_state.listings = []
if "cart" not in st.session_state: st.session_state.cart = {}
if "export_file" not in st.session_state: st.session_state.export_file = None

history = load_history()

st.sidebar.title("🛒 Warenkorb")
st.sidebar.write(f"Ausgewählt: **{len(st.session_state.cart)}**")

# CLOUD FIX: Zwei-Schritt-Export zur Vermeidung von Memory Leaks
if st.session_state.cart:
    if st.sidebar.button("📄 Word Export generieren", use_container_width=True):
        with st.spinner("Erstelle Word-Dokument..."):
            st.session_state.export_file = create_export_buffer()
            st.sidebar.success("Fertig! Bereit zum Download.")

    if st.session_state.export_file:
        st.sidebar.download_button(
            label="⬇️ Datei herunterladen",
            data=st.session_state.export_file,
            file_name=f"Immo_Suche_{datetime.now().strftime('%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

if st.sidebar.button("🗑️ Warenkorb leeren", use_container_width=True): 
    st.session_state.cart = {}
    st.session_state.export_file = None
    st.rerun()

st.sidebar.divider()
st.sidebar.title("⚙️ Einstellungen")
# CLOUD FIX: Standardmäßig True (Headless), da Cloud-Server keine Bildschirme haben
headless_mode = st.sidebar.checkbox("Headless Modus", value=True)
all_srcs = ["Homebase", "Kensington", "TeamPower", "Immowerk", "Robert C. Spies", "Pump Immobilien"]
sel_srcs = st.sidebar.multiselect("Quellen", options=all_srcs, default=all_srcs)

st.title("🏡 Immo Scraper 2026")
col_a, col_b = st.columns([3, 1])
with col_a: t_url = st.text_input("Einzelne URL scrapen", placeholder="https://...")
with col_b: st.write("##"); s_btn = st.button("🚀 SUCHE STARTEN", use_container_width=True)

if s_btn:
    core = ScraperCore(); st.session_state.listings = []
    with st.spinner("Scraping läuft..."):
        if t_url: results = asyncio.run(core.scrape_site(t_url, headless=headless_mode))
        else: results = asyncio.run(core.scrape_all(selected_sources=sel_srcs, headless=headless_mode))
        st.session_state.listings = results
        counts = {}
        for r in results: counts[r.source] = counts.get(r.source, 0) + 1
        if counts:
            cols = st.columns(len(counts))
            for idx, (src, count) in enumerate(counts.items()): cols[idx].metric(src, f"{count} Objekte")
    asyncio.run(core.close_browser())

for i, item in enumerate(st.session_state.listings):
    link = item.link
    if not link: continue
    with st.container(border=True):
        c1, c2 = st.columns([1, 4])
        with c1:
            if item.image_url: st.image(item.image_url, use_container_width=True)
            else: st.info("Kein Bild")
        with c2:
            st.markdown(f"**{item.name}**")
            st.write(f"Ort: {item.location} | qm: {item.area} | Preis: {item.price}")
            if link in history: st.info(f"Exportiert am: {history[link]}")
            ikey = f"cb_{item.source}_{i}_{abs(hash(link))}"
            st.checkbox("Merken", key=ikey, value=(link in st.session_state.cart), on_change=toggle_item, args=(link, item))